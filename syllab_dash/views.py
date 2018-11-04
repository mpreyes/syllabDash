from __future__ import print_function

from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect
from django.conf import settings
from django.core.cache import cache
from docx.api import Document

from dateutil import parser
from dateutil.parser import parse
from itertools import islice

import os  #get timestamp as key for cache: datetime.datetime.now
import rfc3339      # for date object -> date string
import datetime  #get timestamp as key for cache: datetime.datetime.now
from datetime import datetime
# imports for google api
from datetime import *
from apiclient.discovery import build
from oauth2client.file import Storage
from oauth2client.client import OAuth2WebServerFlow
import time
from django.contrib.staticfiles.templatetags.staticfiles import static
from datetime import datetime
from googleapiclient.discovery import build
from httplib2 import Http
from oauth2client import file, client, tools
# from django.templatetags.static import static
import re





# Create your views here.

def file_tables(f,filename): # returns tables in document
    # print("parse files")
    document_tables = []
    document = Document(f) #currently only supports docx files
    for t in document.tables: #m
        # print("__ table __ ")
        # print(t)
        document_tables.append(t)
    return document_tables

#views

def index(request):
    #insertEvents(request)
    return render(request, 'syllab_dash/index.html')

def about(request):
    return render(request, 'syllab_dash/about.html')

def file_upload(request):
    cache_key = 'user_boo' # needs to be unique
    cache_time = 7200 # time in seconds for cache to be valid, 2 hours
    files_parsed = [] #files parsed
    parsed_tables = [] #extracting tables from document
    candidate_tables = [] # tables that contain "date", or "week", or
    parsed_table_data = [] #each row is dictionary, headers mapped to column data
    parsed_assignments = [] #only event data pulled from parsed_table_data
    if request.method == 'POST':
        for f in request.FILES.getlist('file'):
            filename = f.name
            print(filename)
            parsed_tables = file_tables(f,filename)
            candidate_tables = get_tables_cont_dates(parsed_tables)
            parsed_table_data =  parse_table_data(candidate_tables)
            display_table_files = (filename, parsed_table_data)

            parsed_assignments = parse_assignments(parsed_table_data, f)
            parsed_assignments = remove_dates_with_no_assignment(parsed_assignments)
            files_parsed.append(display_table_files)
<<<<<<< HEAD

=======
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
        cache.set(cache_key,files_parsed,cache_time)
        return redirect('list_assignments') #TODO: create a fail page
        #return list_assignments(render,parsed_files_list = files_parsed)
    return render(request, 'syllab_dash/file_upload.html') #TODO: create a fail page


def show_file_contents(request):
    return render(request, 'syllab_dash/show_file_contents.html')



def get_tables_cont_dates(tables): #parse tables, return those containing "date", "week"
    candidate_tables = []
    # print("parse tables, return those containing 'date', 'week'")
    for i in tables:
        for row in i.rows:
            for cell in row.cells:
                if("date" in cell.text.lower() or "week" in cell.text.lower()):
                    candidate_tables.append(i)
                    break
    return candidate_tables

def parse_table_data(candidate_tables):
    data = []
    for c in candidate_tables:
        # print(c)
        for i, row in enumerate(c.rows):

            text = (cell.text for cell in row.cells)
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                print(text)
                keys = tuple(text)
                continue
            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)

    lower_data = []
    for i in data:
        lower_dict = dict((k.lower(), v.strip('\n')) for k, v in i.items())
        lower_data.append(lower_dict)

    return lower_data


def parse_assignments(table_data, file):
    assignments = []
    summary = ""
    for row in table_data:
        #date = parse_date(row["date"])
        for key in row.keys():
            if 'assignment' in key:
                summary = parse_summary(row[key], file)
                break
        timezone = datetime.utcnow().astimezone().tzinfo
        event = {
            'summary': summary,
            'location': '',
            'description': '',
            'start': {
                'dateTime': '2017 11 11',
                'timeZone': timezone,
            },
            'end': {
                'dateTime': '2017 11 11',
                'timeZone': timezone,
            },
            'recurrence': [],
            'attendees': [],
            'reminders': {
                'useDefault': False,
                'overrides': [
                  {'method': 'popup', 'minutes': 24 * 60 * 7},
                  {'method': 'popup', 'minutes': 24 * 60},
                ],
            },
            }
        assignments.append(event)
<<<<<<< HEAD

=======
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
    return assignments

def parse_summary(assignment, file):
    document = Document(file)
    regex = re.compile("[A-Z][A-Z][A-Z] [0-9][0-9][0-9][0-9]")
    course_title = ''
    for paragraph in document.paragraphs:
        result = regex.match(paragraph.text)
        if(result):
            course_title = result.group()
            break
    if assignment:
        assignment = assignment.replace('\n', ', ')
        summary = course_title + ' ' + assignment
    else:
        summary = 'N/A'
    return summary


def parse_date(date):
    date = parser.parse(date, fuzzy=True)
    if date.year:
        datetime_object = rfc3339.rfc3339(date) #change to rfc3339 format
    else:
        datetime_object = rfc3339.rfc3339(date) #change to rfc3339 format
<<<<<<< HEAD

=======
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
    print(datetime_object)
    return datetime_object

def remove_dates_with_no_assignment(assignments):
    for i, entry in enumerate(assignments):
        if entry['summary'] == 'N/A':
            assignments.pop(i)
    return assignments

def list_assignments(request):
    data = cache.get("user_boo")
    print(data)
    for i in data: #here
        if i != None:
            filename, table = i
            print(filename)
            print(table)

<<<<<<< HEAD
    insertEvents()

=======
    #insertEvents()
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
    # for i in data:
    #     with i.open() as f:
    #         document = Document(f) #currently only supports docx files
    #         for p in document.paragraphs: #maybe split up file detection into separate functions
    #             #print(p.text)  # parsePdf(p), parseWord(p), etc ?
    #             parsed_files.append(p.text)
    #             #break


        #print(data)
    return render(request, 'syllab_dash/list_assignments.html',{"file_data": data })


def finished_upload(request):
    return render(request, 'syllab_dash/finished_upload.html')





def insertEvents():
    SCOPES = 'https://www.googleapis.com/auth/calendar'

<<<<<<< HEAD
    store = file.Storage('token.json')
    creds = store.get()
=======
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
    event = {
    'summary': 'NEW SYLLAB DASH EVENT -- HI',
    'location': '800 Howard St., San Francisco, CA 94103',
    'description': 'A chance to hear more about Google\'s developer products.',
    'start': {
        'dateTime': '2019-01-02T09:00:00-07:00',
        'timeZone': 'America/Chicago',
    },
    'end': {
        'dateTime': '2019-01-02T17:00:00-07:00',
        'timeZone': 'America/Chicago',
    },
    'reminders': {
        'useDefault': False,
        'overrides': [
        {'method': 'email', 'minutes': 24 * 60},
        {'method': 'popup', 'minutes': 10},
        ],
    },
    }


    store = file.Storage('token.json')
    creds = store.get()
    if not creds or creds.invalid:
        flow = client.flow_from_clientsecrets('credentials.json', SCOPES)
        creds = tools.run_flow(flow, store)
    service = build('calendar', 'v3', http=creds.authorize(Http()))

<<<<<<< HEAD
    # calendar_list_entry = service.calendarList().get(calendarId= 'primary' ).execute()
    # print(calendar_list_entry['summary'])


    # Call the Calendar API
    # now = datetime.utcnow().isoformat() + 'Z' # 'Z' indicates UTC time
    # acl = service.acl().list(calendarId='primary').execute()

    # for rule in acl['items']:
    #     print (rule['id'])
    #     print(rule['role'])

    event = service.events().insert(calendarId = 'primary', body=event).execute()
    print('Event created:')
    print(event.get('htmlLink'))

=======
    # Call the Calendar API
    now = datetime.utcnow().isoformat() + 'Z' # 'Z' indicates UTC time
    print('Getting the upcoming 10 events')
    myCalendarId = "reyes.madelyn.mr@gmail.com"
    events_result = service.events().list(calendarId= myCalendarId, timeMin=now,
                                        maxResults=10, singleEvents=True,
                                        orderBy='startTime').execute()
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
    #events = events_result.get('items', [])

    # if not events:
    #     print('No upcoming events found.')
    # for event in events:
    #     start = event['start'].get('dateTime', event['start'].get('date'))
    #     print(start, event['summary'])

    event = service.events().insert(calendarId='primary', body=event).execute()

    calendar_list_entry = service.calendarList().get(calendarId='primary').execute()
    print(calendar_list_entry['summary'])

    print('Event created:')
    print(event.get('htmlLink'))

<<<<<<< HEAD
    # SCOPES = 'https://www.googleapis.com/auth/calendar'

    # event = {

    #   'summary': 'This is a test event summary.',
    #   'start': {
    #     'dateTime': '2018-11-02T22:10:00',
    #     'timeZone': time.tzname[time.daylight]
    #   },
    #   'end': {
    #     'dateTime': '2018-11-02T23:00:00',
    #     'timeZone': time.tzname[time.daylight]
    #   },
    #   'reminders': {
    #     'useDefault': False,
    #     'overrides': [
    #       {'method': 'email', 'days': 7},
    #       {'method': 'popup', 'minutes': 10},
    #     ],
    #   },

    # }
    # event = service.events().insert(calendarId='primary', body=event).execute()
    # print('Event created:')
    # print(event.get('htmlLink'))


#     testEvent2 = {
#       'summary': 'This is a test event summary.',
#       'start': {
#         'dateTime': '2019-11-09T22:10:00',
#         'timeZone': time.tzname[time.daylight]
#       },
#       'end': {
#         'dateTime': '2018-11-09T23:00:00',
#         'timeZone': time.tzname[time.daylight]
#       },
#       'reminders': {
#         'useDefault': False,
#         'overrides': [
#           {'method': 'email', 'days': 7},
#           {'method': 'popup', 'minutes': 10},
#         ],
#       },
#     }




#    # Shows basic usage of the Google Calendar API.
#    # Prints the start and name of the next 10 events on the user's calendar.


#     flow = OAuth2WebServerFlow(
#         client_id='319052537199-fndghhjj6akqht9gmooe818k5b00jnp6.apps.googleusercontent.com',
#         client_secret='6yYGMakT8_lBX4mTiUr7yfb5',
#         scope='https://www.googleapis.com/auth/calendar',
#         user_agent='Syllab-Dash',
#     )
#     storage = Storage('calendar.dat')
#     credentials = storage.get()

#     code = request.GET.get('code')
#     if credentials is None or credentials.invalid == True:
#         oauth_callback = 'index.html'
#         flow.redirect_uri = oauth_callback
#         flow.step1_get_authorize_url()
#         credential = flow.step2_exchange(code, http=None)
#         storage.put(credential)
#         credential.set_store(storage)
#     http = httplib2.Http()
#     http = credentials.authorize(http)

#     service = build(serviceName='calendar', version='v3', http=http,
#                     developerKey='AIzaSyBP60OCOPNIXTWVHG-XmorCqvBsjzThdFQ')

#     event = service.events().insert(calendarId='primary', body=testEvent).execute()
#     event2 = service.events().insert(calendarId='primary', body=testEvent2).execute()

#     if not event:
#         print("Error with adding event 1")
#     else:
#         print("Added event 1 successfully... maybe.")
#     if not event2:
#         print("Error with adding event 2")
#     else:
#         print("Added event 2 successfully... maybe.")
=======
>>>>>>> fe16e520ddc206ddc7e8fa251419492333604f9c
